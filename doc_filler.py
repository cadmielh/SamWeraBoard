"""Fill Word (.docx) and Google Docs templates with extracted ID data."""

import copy
import io
import re
from docx import Document
from docx.text.paragraph import Paragraph


def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    """Replace placeholders in a paragraph, handling runs that split a placeholder."""
    # Rebuild full paragraph text from runs
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text
    for placeholder, value in replacements.items():
        new_text = new_text.replace(placeholder, value)

    if new_text == full_text:
        return  # Nothing changed

    # Put the whole replaced text into the first run, clear the rest
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def _para_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs).strip()


def _expand_repeat_blocks(doc: Document, groups: dict[str, list[dict[str, str]]]) -> None:
    """
    Expand {{#TAG}} ... {{/TAG}} paragraph ranges (top-level body paragraphs) into
    one copy of the enclosed paragraphs per item in groups[TAG], substituting
    singular placeholders (e.g. {{NUME}}, {{CNP}}, {{INDEX}}) from that item.
    The original marker paragraphs and the template block are removed afterwards.
    Tags with no matching group, or malformed (missing end tag), are left as-is.
    """
    for tag, items in groups.items():
        start_marker = f"{{{{#{tag}}}}}"
        end_marker = f"{{{{/{tag}}}}}"

        while True:
            paragraphs = doc.paragraphs
            start_idx = next((i for i, p in enumerate(paragraphs) if _para_text(p) == start_marker), None)
            if start_idx is None:
                break
            end_idx = next(
                (i for i in range(start_idx + 1, len(paragraphs)) if _para_text(paragraphs[i]) == end_marker),
                None,
            )
            if end_idx is None:
                break  # no matching end tag — leave the stray start marker as-is

            block_paragraphs = paragraphs[start_idx + 1:end_idx]
            anchor = paragraphs[end_idx]._p

            for i, item in enumerate(items, start=1):
                person = {**item, "INDEX": str(i)}
                item_replacements = {"{{" + k + "}}": v for k, v in person.items()}
                for bp in block_paragraphs:
                    clone = copy.deepcopy(bp._p)
                    anchor.addprevious(clone)
                    _replace_in_paragraph(Paragraph(clone, bp._parent), item_replacements)

            # Remove the original template block + both markers
            for bp in block_paragraphs:
                bp._p.getparent().remove(bp._p)
            paragraphs[end_idx]._p.getparent().remove(paragraphs[end_idx]._p)
            paragraphs[start_idx]._p.getparent().remove(paragraphs[start_idx]._p)


_NUMBERED_TAG_RE = re.compile(r"\{\{(ASOCIAT|ADMINISTRATOR|MEMBRU_IF)_(\d+)_[A-Z_]+\}\}")
_NUMBERED_KEY_RE = re.compile(r"^\{\{(ASOCIAT|ADMINISTRATOR|MEMBRU_IF)_(\d+)_")


def _max_numbered_index(replacements: dict[str, str]) -> dict[str, int]:
    """Highest N actually present for each known numbered prefix (ASOCIAT_N_*, ...)."""
    max_idx: dict[str, int] = {}
    for key in replacements:
        m = _NUMBERED_KEY_RE.match(key)
        if m:
            prefix, n = m.group(1), int(m.group(2))
            if n > max_idx.get(prefix, 0):
                max_idx[prefix] = n
    return max_idx


def _has_unresolved_numbered_tag(text: str, max_idx: dict[str, int]) -> bool:
    for m in _NUMBERED_TAG_RE.finditer(text):
        prefix, n = m.group(1), int(m.group(2))
        if n > max_idx.get(prefix, 0):
            return True
    return False


def _clean_unresolved_numbered_positions(doc: Document, replacements: dict[str, str]) -> None:
    """
    Known numbered tags (ASOCIAT_N_*, ADMINISTRATOR_N_*, MEMBRU_IF_N_*) that
    reference a position beyond how many actually exist — e.g. {{ASOCIAT_3_NUME}}
    when there are only 2 asociați — are cleaned up instead of left as literal
    braces in the output:
      - in the main body/headers/footers, the whole paragraph is deleted (a
        half-finished sentence reads worse than no sentence);
      - in a table cell, only the cell's text is cleared, so the table itself
        isn't structurally broken.
    Unrecognized/misspelled tags are left untouched, so they stay visible —
    a clear signal something needs fixing before the document is final.
    """
    max_idx = _max_numbered_index(replacements)
    if not max_idx:
        return

    def delete_paragraph(paragraph) -> None:
        p = paragraph._p
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)

    for paragraph in list(doc.paragraphs):
        text = "".join(run.text for run in paragraph.runs)
        if _has_unresolved_numbered_tag(text, max_idx):
            delete_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = "".join(run.text for run in paragraph.runs)
                    if _has_unresolved_numbered_tag(text, max_idx):
                        for run in paragraph.runs:
                            run.text = ""

    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header is not None:
                for paragraph in list(header.paragraphs):
                    text = "".join(run.text for run in paragraph.runs)
                    if _has_unresolved_numbered_tag(text, max_idx):
                        delete_paragraph(paragraph)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer is not None:
                for paragraph in list(footer.paragraphs):
                    text = "".join(run.text for run in paragraph.runs)
                    if _has_unresolved_numbered_tag(text, max_idx):
                        delete_paragraph(paragraph)


def fill_docx(
    template_bytes: bytes,
    replacements: dict[str, str],
    groups: dict[str, list[dict[str, str]]] | None = None,
) -> bytes:
    """
    Fill a .docx template by replacing {{PLACEHOLDER}} markers.

    If `groups` is given, {{#TAG}}...{{/TAG}} blocks are expanded first — once
    per item in groups[TAG] — before the flat placeholder pass runs over the
    whole (now expanded) document.

    Returns the filled document as bytes.
    """
    doc = Document(io.BytesIO(template_bytes))

    if groups:
        _expand_repeat_blocks(doc, groups)

    # Replace in main body paragraphs
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)

    # Replace in headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header is not None:
                for paragraph in header.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer is not None:
                for paragraph in footer.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)

    _clean_unresolved_numbered_positions(doc, replacements)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def list_placeholders_in_docx(template_bytes: bytes) -> list[str]:
    """Return all unique {{PLACEHOLDER}} markers found in the template.

    Excludes {{#TAG}}/{{/TAG}} repeat-block markers — those are structural,
    not fillable fields.
    """
    doc = Document(io.BytesIO(template_bytes))
    pattern = re.compile(r"\{\{[^}]+\}\}")
    marker_pattern = re.compile(r"^\{\{[#/]")
    found: set[str] = set()

    def scan_paragraphs(paragraphs):
        for p in paragraphs:
            text = "".join(r.text for r in p.runs)
            for match in pattern.findall(text):
                if not marker_pattern.match(match):
                    found.add(match)

    scan_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                scan_paragraphs(cell.paragraphs)

    return sorted(found)
